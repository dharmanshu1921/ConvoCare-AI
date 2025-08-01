import streamlit as st
import uuid
import google.generativeai as genai
from geopy.geocoders import Nominatim
import folium
from streamlit_folium import st_folium
from streamlit_js_eval import get_geolocation
import os
from dotenv import load_dotenv
from PIL import Image
import PyPDF2
import docx
from pdf2image import convert_from_bytes
import magic
from prompts import primary_prompt, plan_prompt, sim_prompt, num_prompt, policy_prompt, faq_prompt
from tools import plan_tools, sim_tools, num_tools, policy_tools, faq_tools
from agents import run_chatbot, State, HumanMessage, graph, memory, listen_to_user_whisper
from langchain_core.messages import AIMessage, ToolMessage
import speech_recognition as sr
from io import BytesIO
import json
import re
import pyttsx3

# Load environment variables
load_dotenv()

# Streamlit page configuration
st.set_page_config(page_title="Airtel Customer Support & File Analysis", layout="wide")

# Initialize session state
if "thread_id" not in st.session_state:
    st.session_state.thread_id = str(uuid.uuid4())
if "messages" not in st.session_state:
    st.session_state.messages = []  # For chatbot tab
if "analysis_messages" not in st.session_state:
    st.session_state.analysis_messages = []  # For file analysis tab
if "language" not in st.session_state:
    st.session_state.language = "en"
if 'search_state' not in st.session_state:
    st.session_state.search_state = "init"  # init, searching, results, error
    st.session_state.stores = []
    st.session_state.current_location = None
    st.session_state.radius = 5
    st.session_state.gemini_api_key = os.getenv("GEMINI_API_KEY")
    st.session_state.gemini_configured = False
    st.session_state.gemini_model = None
    st.session_state.voice_input = ""
    st.session_state.recording = False

# Gemini API Configuration
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    st.error("GEMINI_API_KEY environment variable not set. Create a .env file or set it in your shell before running the app.")
    st.stop()
genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-2.0-flash')

# Configure Gemini API for store locator
if st.session_state.gemini_api_key and not st.session_state.gemini_configured:
    try:
        genai.configure(api_key=st.session_state.gemini_api_key)
        st.session_state.gemini_configured = True
        st.session_state.gemini_model = genai.GenerativeModel('gemini-2.0-flash')
    except Exception as e:
        st.error(f"Error configuring Gemini: {e}")

# Initialize geolocator for store locator
geolocator = Nominatim(user_agent="airtel_store_locator")

# Updated text-to-speech function using pyttsx3
def speak_output(text):
    """Generate and return audio bytes using pyttsx3"""
    if not text or not isinstance(text, str):
        st.error("No valid text provided for audio generation.")
        return None
    
    # Clean text to remove problematic characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
    text = text.strip()
    if not text:
        st.error("Text is empty after cleaning.")
        return None

    try:
        # Initialize pyttsx3 engine
        engine = pyttsx3.init()
        
        # Set language properties (approximate mapping for en, es, fr, etc.)
        if st.session_state.language == "en":
            engine.setProperty('voice', 'com.apple.speech.synthesis.voice.Alex' if os.name == 'posix' else 'HKEY_LOCAL_MACHINE\SOFTWARE\Microsoft\Speech\Voices\Tokens\TTS_MS_EN-US_DAVID_11.0')
        # Add more language mappings if needed
        engine.setProperty('rate', 150)  # Adjust speech rate if needed
        
        # Save audio to BytesIO
        audio_bytes = BytesIO()
        engine.save_to_file(text[:500], 'temp.wav')  # Save to temporary file
        engine.runAndWait()
        
        # Read the temporary file into BytesIO
        with open('temp.wav', 'rb') as f:
            audio_bytes.write(f.read())
        audio_bytes.seek(0)
        
        # Verify audio data
        audio_data = audio_bytes.read()
        if len(audio_data) == 0:
            st.error("Generated audio is empty.")
            return None
        audio_bytes.seek(0)
        
        # Clean up temporary file
        if os.path.exists('temp.wav'):
            os.remove('temp.wav')
        
        return audio_bytes
    except Exception as e:
        st.error(f"Error generating audio: {str(e)}")
        return None

# Functions for File Analysis
def get_file_type(uploaded_file):
    """Determine file type using magic numbers"""
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(uploaded_file.getvalue())
    return file_type

def display_file(uploaded_file, file_type):
    """Display uploaded file based on its type"""
    if 'image' in file_type:
        st.image(uploaded_file, caption="Uploaded Image", use_container_width=True)
    elif file_type == 'application/pdf':
        images = convert_from_bytes(uploaded_file.getvalue(), first_page=1, last_page=1)
        st.image(images[0], caption="First Page Preview", use_container_width=True)
    elif file_type in ['text/plain', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        st.info("Document Content Preview:")
        if file_type == 'text/plain':
            content = uploaded_file.getvalue().decode("utf-8")
            st.text_area("", value=content, height=200)
        else:  # DOCX
            doc = docx.Document(uploaded_file)
            full_text = [paragraph.text for paragraph in doc.paragraphs]
            st.text_area("", value="\n".join(full_text), height=200)

def extract_content(uploaded_file, file_type):
    """Extract content based on file type"""
    content = ""
    if 'image' in file_type:
        return Image.open(uploaded_file)
    elif file_type == 'application/pdf':
        reader = PyPDF2.PdfReader(uploaded_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                content += page_text + "\n"
        return content
    elif file_type == 'text/plain':
        return uploaded_file.getvalue().decode("utf-8")
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    return None

def generate_explanation(content, user_query=None, is_image=False):
    """Generate explanation using Gemini Flash, incorporating user query if provided"""
    prompt = (
        "Analyze this content comprehensively. Explain key elements, summarize main points, and provide insights. Focus on:\n"
        "- Key themes and concepts\n"
        "- Important details\n"
        "- Overall context and significance\n"
    )
    if user_query:
        prompt += f"\n\nUser Query: {user_query}\nPlease incorporate the user's query into the analysis where relevant."
    prompt += "\nProvide your response in clear, structured markdown format."
    try:
        if is_image:
            response = model.generate_content([prompt, content])
        else:
            response = model.generate_content(prompt + (f"\n\nContent:\n{content}" if content else ""))
        return response.text
    except Exception as e:
        return f"Error generating explanation: {str(e)}"

# Robust JSON extraction function for store locator
def extract_json_from_text(text):
    """Extract JSON content from Gemini's response"""
    try:
        clean_text = text.strip()
        clean_text = re.sub(r'```json|```', '', clean_text)
        start_idx = clean_text.find('{')
        end_idx = clean_text.rfind('}') + 1
        if start_idx == -1 or end_idx == 0:
            return None
        json_str = clean_text[start_idx:end_idx]
        return json.loads(json_str)
    except Exception as e:
        st.error(f"JSON extraction error: {e}")
        return None

# Function to call Gemini and get Airtel stores
def get_airtel_stores(lat, lon, radius):
    if not st.session_state.gemini_model:
        return []

    prompt = f"""
    Find all Airtel stores within a {radius} km radius of these coordinates: 
    Latitude: {lat}, Longitude: {lon}.
    
    Return ONLY a JSON object with the following structure:
    {{
        "stores": [
            {{
                "name": "Store name",
                "address": "Full store address",
                "latitude": 12.3456,
                "longitude": 77.6789,
                "distance_km": 1.2
            }}
        ]
    }}
    
    Important:
    1. Return ONLY the JSON object with no additional text
    2. Include exact latitude/longitude coordinates
    3. Calculate distance in km
    4. Return at least 5 stores if available
    """
    try:
        response = st.session_state.gemini_model.generate_content(prompt)
        json_text = "".join(part.text for part in response.candidates[0].content.parts)
        data = extract_json_from_text(json_text)
        if not data:
            return []

        return data.get("stores", [])
    except Exception as e:
        st.error(f"⚠️ Gemini API error: {e}")
        return []

# Convert address to coordinates
def get_coordinates_from_address(address):
    try:
        location = geolocator.geocode(address)
        return (location.latitude, location.longitude) if location else None
    except Exception as e:
        st.error(f"Geocoding error: {e}")
        return None

# even though I don't need this for audio, I'm keeping it for the store locator
def recognize_speech():
    r = sr.Recognizer()
    with st.spinner("Listening... Speak now"):
        with sr.Microphone() as source:
            audio = r.listen(source)
        try:
            text = r.recognize_google(audio)
            return text
        except sr.UnknownValueError:
            st.error("Could not understand audio")
        except sr.RequestError as e:
            st.error(f"Could not request results: {e}")
    return None

# Create three tabs
tab1, tab2, tab3 = st.tabs(["ConvoCare AI", "Document & Image Analysis", "Store Locator"])

# Tab 1: Airtel Customer Support Chatbot
with tab1:
    st.title("🤖 ConvoCare AI")
    st.markdown("""
    Welcome to the ConvoCare AI! Ask about plans, SIM issues, customer care numbers, policies, FAQs, or store locations. 
    Type or speak your query below, or resume a previous conversation using a thread ID. 
    You can also choose to hear the assistant's response.
    """)

    # Sidebar for thread ID and language selection
    with st.sidebar:
        st.header("Chat Settings")
        thread_id_input = st.text_input("Enter Thread ID (or leave blank for new session):",
                                        value=st.session_state.thread_id)
        if thread_id_input != st.session_state.thread_id:
            st.session_state.thread_id = thread_id_input or str(uuid.uuid4())
            st.session_state.messages = []
            st.rerun()

        if st.button("Clear Conversation"):
            st.session_state.messages = []
            st.session_state.thread_id = str(uuid.uuid4())
            st.rerun()

    # Chat interface
    st.subheader("Chat with Airtel Support")
    chat_container = st.container()

    # Display conversation history
    with chat_container:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    # User input (text or speech)
    col1, col2 = st.columns([4, 1])
    with col1:
        user_input = st.chat_input("Type your query (e.g., 'Best plan for 5GB/day')")
    with col2:
        if st.button("🎤 Speak Query"):
            with st.spinner("🎤 Listening... Speak now (5 seconds):"):
                try:
                    spoken_input = listen_to_user_whisper(duration=5, language=st.session_state.language)
                    if spoken_input:
                        user_input = spoken_input
                        st.success(f"🗣️ You said: {spoken_input}")
                    else:
                        st.error("❌ Could not transcribe speech. Please try again or type your query.")
                except Exception as e:
                    st.error(f"❌ Speech recognition error: {e}")

    # Process user input (typed or spoken)
    if user_input:
        # Add user message to history
        st.session_state.messages.append({"role": "user", "content": user_input})

        # Display user message
        with chat_container:
            with st.chat_message("user"):
                st.markdown(user_input)

        # Run chatbot with the current thread ID and user input
        try:
            with st.spinner("Processing..."):
                response = run_chatbot(user_input, st.session_state.thread_id, st.session_state.language)

            if response:
                # Add assistant response to history
                st.session_state.messages.append({"role": "assistant", "content": response})

                # Display assistant response
                with chat_container:
                    with st.chat_message("assistant"):
                        st.markdown(response)

                        # Option to hear the response
                        # if st.button("🔊 Hear Response", key=f"audio_{len(st.session_state.messages)}"):
                            # st.write(f"Debug: Response text = '{response}'")  # Debug log
                            # audio_bytes = speak_output(response)
                            # if audio_bytes:
                            #     st.audio(audio_bytes, format="audio/wav")
                            # else:
                            #     st.error("Failed to generate audio.")

            else:
                st.error("No response from the assistant. Please try again.")

        except Exception as e:
            st.error(f"Error processing request: {e}")

    # Show conversation history
    if st.button("Show Full Conversation History"):
        checkpoint = memory.get({"configurable": {"thread_id": st.session_state.thread_id}})
        if checkpoint and "messages" in checkpoint:
            st.subheader("Full Conversation History")
            for msg in checkpoint["messages"]:
                role = msg.type
                content = msg.content
                with st.chat_message(role):
                    st.markdown(f"**{role.capitalize()}**: {content}")
        else:
            st.info("No conversation history available.")

# Tab 2: Document & Image Analysis
with tab2:
    st.title("📄 Document & Image Explainer")
    st.markdown("Upload an image or document (PDF, DOCX, TXT) and optionally provide a query to guide the AI-powered analysis.")

    # File uploader
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=["png", "jpg", "jpeg", "pdf", "docx", "txt"]
    )

    # Analysis chat interface
    st.subheader("Analysis Query")
    analysis_container = st.container()

    # Display analysis history
    with analysis_container:
     for message in st.session_state.analysis_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])


    # User input for analysis (text or speech)
    col1, col2 = st.columns([4, 1])
    with col1:
        analysis_input = st.chat_input("Type your query for analysis (e.g., 'Summarize this document' or 'Extract key points from this image')")
    with col2:
        if st.button("🎤 Speak Analysis Query"):
            with st.spinner("🎤 Listening... Speak now (5 seconds):"):
                try:
                    spoken_input = listen_to_user_whisper(duration=5, language=st.session_state.language)
                    if spoken_input:
                        analysis_input = spoken_input
                        st.success(f"🗣️ You said: {spoken_input}")
                    else:
                        st.error("❌ Could not transcribe speech. Please try again or type your query.")
                except Exception as e:
                    st.error(f"❌ Speech recognition error: {e}")

    # Process file and user query
    if uploaded_file or analysis_input:
        if uploaded_file:
            file_type = get_file_type(uploaded_file)
            st.subheader("Uploaded File Preview")
            display_file(uploaded_file, file_type)

        if analysis_input:
            # Add user query to analysis history
            st.session_state.analysis_messages.append({"role": "user", "content": analysis_input})

            # Display user query
            with analysis_container:
                with st.chat_message("user"):
                    st.markdown(analysis_input)

        st.subheader("Analysis Results")
        with st.spinner("Analyzing content with Gemini 2.0 Flash..."):
            if uploaded_file:
                is_image = 'image' in file_type
                content = extract_content(uploaded_file, file_type)
                if content is not None:
                    explanation = generate_explanation(content, user_query=analysis_input, is_image=is_image)
                    # Add assistant response to analysis history
                    st.session_state.analysis_messages.append({"role": "assistant", "content": explanation})
                    # Display assistant response
                    with analysis_container:
                        with st.chat_message("assistant"):
                            st.markdown(explanation)
                            # Option to hear the response
                            # if st.button("🔊 Hear Response", key=f"analysis_audio_{len(st.session_state.analysis_messages)}"):
                            #     st.write(f"Debug: Explanation text = '{explanation}'")  # Debug log
                            #     audio_bytes = speak_output(explanation)
                            #     if audio_bytes:
                            #         st.audio(audio_bytes, format="audio/wav")
                            #     else:
                            #         st.error("Failed to generate audio.")
                else:
                    st.error("Unsupported file format or extraction error.")
            else:
                # Handle query without file
                explanation = generate_explanation("", user_query=analysis_input)
                st.session_state.analysis_messages.append({"role": "assistant", "content": explanation})
                with analysis_container:
                    with st.chat_message("assistant"):
                        st.markdown(explanation)
                        # if st.button("🔊 Hear Response", key=f"analysis_audio_{len(st.session_state.analysis_messages)}"):
                        #     st.write(f"Debug: Explanation text = '{explanation}'")  # Debug log
                        #     audio_bytes = speak_output(explanation)
                        #     if audio_bytes:
                        #         st.audio(audio_bytes, format="audio/wav")
                        #     else:
                        #         st.error("Failed to generate audio.")

# Tab 3: Store Locator
with tab3:
    st.title("📍 Airtel Store Locator")
    st.markdown("Find nearby Airtel stores using your location, voice command, or by entering an address.")

    # Create two columns for input methods
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📍 Use My Location")
        detect_btn = st.button("Detect My Location", key="detect", use_container_width=True)

    with col2:
        st.subheader("🔍 Search by Address")
        address_col, voice_col = st.columns([4, 1])
        
        with address_col:
            address = st.text_input(
                "Enter address or city:",
                value=st.session_state.voice_input,
                key="address_input",
                label_visibility="collapsed",
                placeholder="Enter address or use voice"
            )
        
        with voice_col:
            voice_btn = st.button("🎤", key="voice_btn", use_container_width=True)
        
        search_btn = st.button("Find Stores", key="search", use_container_width=True)

    # Radius slider
    radius = st.slider("Search radius (km):", 1, 30, st.session_state.radius, key="radius_slider")
    if radius != st.session_state.radius:
        st.session_state.radius = radius
        if st.session_state.search_state == "results":
            st.session_state.search_state = "init"

    # Handle voice input
    if voice_btn:
        st.session_state.recording = True
        st.session_state.voice_input = recognize_speech() or ""
        st.session_state.recording = False
        st.rerun()

    # Handle location detection
    if detect_btn:
        st.session_state.search_state = "searching"
        try:
            location = get_geolocation()
            if location and 'coords' in location:
                st.session_state.current_location = (
                    location['coords']['latitude'],
                    location['coords']['longitude']
                )
                st.success("📍 Location detected successfully!")
            else:
                st.error("Could not detect location. Try again or use address.")
                st.session_state.search_state = "error"
        except Exception as e:
            st.error(f"Location error: {e}")
            st.session_state.search_state = "error"

    # Handle address search
    if search_btn:
        if address:
            st.session_state.search_state = "searching"
            coords = get_coordinates_from_address(address)
            if coords:
                st.session_state.current_location = coords
                st.success(f"📍 Location found: {address}")
            else:
                st.error("Could not find address. Try another.")
                st.session_state.search_state = "error"
        else:
            st.error("Please enter a valid address.")
            st.session_state.search_state = "error"

    # Perform search if needed
    if (st.session_state.search_state == "searching" and 
        st.session_state.current_location and
        st.session_state.gemini_configured):
        
        lat, lon = st.session_state.current_location
        with st.spinner(f"🔍 Searching for Airtel stores within {st.session_state.radius} km..."):
            st.session_state.stores = get_airtel_stores(lat, lon, st.session_state.radius)
            st.session_state.search_state = "results" if st.session_state.stores else "no_results"

    # Display results
    if st.session_state.search_state == "results":
        lat, lon = st.session_state.current_location
        stores = st.session_state.stores
        
        # Create map
        m = folium.Map(location=[lat, lon], zoom_start=12)
        
        # Add user location marker
        folium.Marker(
            [lat, lon], 
            popup="Your Location", 
            icon=folium.Icon(color='blue', icon='user')
        ).add_to(m)
        
        # Add store markers
        for store in stores:
            folium.Marker(
                [store["latitude"], store["longitude"]],
                popup=f"<b>{store['name']}</b><br>{store['address']}",
                icon=folium.Icon(color='red', icon='shopping-cart')
            ).add_to(m)
        
        # Fit map to show all markers
        if stores:
            bounds = [[lat, lon]] + [[s["latitude"], s["longitude"]] for s in stores]
            m.fit_bounds(bounds)
        
        # Display map
        st_folium(m, width=1200, height=500)

        # Show store list
        st.subheader(f"📌 Found {len(stores)} Airtel stores:")
        for i, store in enumerate(stores, 1):
            gmap_link = f"https://www.google.com/maps/search/?api=1&query={store['latitude']},{store['longitude']}"
            st.markdown(f"""
            <div style="border:1px solid #e0e0e0; border-radius:10px; padding:15px; margin-bottom:15px;">
                <h4>{i}. {store['name']}</h4>
                <p>📍 <b>Address:</b> {store['address']}</p>
                <p>📏 <b>Distance:</b> {store['distance_km']:.1f} km</p>
                <a href="{gmap_link}" target="_blank" style="text-decoration:none;">
                    <button style="background-color:#4285F4; color:white; border:none; padding:8px 15px; border-radius:4px; cursor:pointer;">
                        🗺️ View on Google Maps
                    </button>
                </a>
            </div>
            """, unsafe_allow_html=True)
        
        # Add reset button
        if st.button("New Search"):
            st.session_state.search_state = "init"
            st.session_state.current_location = None
            st.session_state.stores = []
            st.session_state.voice_input = ""
            st.rerun()

    elif st.session_state.search_state == "no_results":
        st.warning("No Airtel stores found. Try a different location or increase the radius.")
        if st.button("Try Again"):
            st.session_state.search_state = "init"
            st.rerun()

    elif st.session_state.search_state == "error":
        if st.button("Retry Search"):
            st.session_state.search_state = "init"
            st.rerun()

    elif st.session_state.current_location and not st.session_state.gemini_configured:
        st.error("Gemini API not configured. Please set GEMINI_API_KEY in your .env file.")

    elif st.session_state.search_state == "init":
        st.info("👋 Start by detecting your location, using voice, or entering an address.")

    # Show recording indicator
    if st.session_state.recording:
        st.warning("Recording in progress... Please speak now")

st.markdown("---")