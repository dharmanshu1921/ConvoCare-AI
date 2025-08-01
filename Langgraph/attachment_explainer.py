# app.py

import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
from PIL import Image
import PyPDF2
import docx
from pdf2image import convert_from_bytes
import magic

# Load environment variables from .env (if exists)
load_dotenv()

# ------------ Gemini API Configuration ------------
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    st.error("GEMINI_API_KEY environment variable not set. Create a .env file or set it in your shell before running the app.")
    st.stop()
genai.configure(api_key=api_key)

# Initialize Gemini 2.0 Flash model
model = genai.GenerativeModel('gemini-2.0-flash')


def get_file_type(uploaded_file):
    """Determine file type using magic numbers"""
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(uploaded_file.getvalue())
    return file_type

def display_file(uploaded_file, file_type):
    """Display uploaded file based on its type"""
    if 'image' in file_type:
        st.image(uploaded_file, caption="Uploaded Image", use_container_width=True)
    elif file_type == 'application/pdf':
        images = convert_from_bytes(uploaded_file.getvalue(), first_page=1, last_page=1)
        st.image(images[0], caption="First Page Preview", use_container_width=True)
    elif file_type in ['text/plain', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        st.info("Document Content Preview:")
        if file_type == 'text/plain':
            content = uploaded_file.getvalue().decode("utf-8")
            st.text_area("", value=content, height=200)
        else:  # DOCX
            doc = docx.Document(uploaded_file)
            full_text = [paragraph.text for paragraph in doc.paragraphs]
            st.text_area("", value="\n".join(full_text), height=200)

def extract_content(uploaded_file, file_type):
    """Extract content based on file type"""
    content = ""
    if 'image' in file_type:
        return Image.open(uploaded_file)
    elif file_type == 'application/pdf':
        reader = PyPDF2.PdfReader(uploaded_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                content += page_text + "\n"
        return content
    elif file_type == 'text/plain':
        return uploaded_file.getvalue().decode("utf-8")
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    return None

def generate_explanation(content, is_image=False):
    """Generate explanation using Gemini Flash"""
    prompt = (
        "Analyze this content comprehensively. Explain key elements, summarize main points, and provide insights. Focus on:\n"
        "- Key themes and concepts\n"
        "- Important details\n"
        "- Overall context and significance\n"
        "Provide your response in clear, structured markdown format."
    )
    try:
        if is_image:
            response = model.generate_content([prompt, content])
        else:
            response = model.generate_content(prompt + "\n\n" + content)
        return response.text
    except Exception as e:
        return f"Error generating explanation: {str(e)}"

# ----------------- Streamlit UI ------------------
st.title("📄 Gemini Document & Image Explainer")
st.markdown("Upload any image or document (PDF, DOCX, TXT) for AI-powered analysis.")

uploaded_file = st.file_uploader(
    "Choose a file",
    type=["png", "jpg", "jpeg", "pdf", "docx", "txt"]
)

if uploaded_file:
    file_type = get_file_type(uploaded_file)
    st.subheader("Uploaded File Preview")
    display_file(uploaded_file, file_type)

    st.subheader("Analysis Results")
    with st.spinner("Analyzing content with Gemini 2.0 Flash..."):
        is_image = 'image' in file_type
        content = extract_content(uploaded_file, file_type)
        if content is not None:
            explanation = generate_explanation(content, is_image)
            st.markdown(explanation)
        else:
            st.error("Unsupported file format or extraction error.")

st.markdown("---")
